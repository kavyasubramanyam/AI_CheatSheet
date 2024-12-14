from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
from pylatexenc.latex2text import LatexNodes2Text

def clean_text(text):
    #clean text from latex to json
    if text is None:
        return ""
    
    latex_converter = LatexNodes2Text(math_mode='text', 
                                    keep_braced_groups=False)
    cleaned = latex_converter.latex_to_text(str(text))
    
    # Remove any extra quotes or brackets
    cleaned = cleaned.replace('[', '').replace(']', '')
    cleaned = cleaned.replace('\'', '').replace('\"', '')
    
    # line breaks and spacing
    lines = cleaned.split('\n')
    cleaned_lines = [line.strip() for line in lines if line.strip()]
    result = '\n'.join(cleaned_lines)
    
    return result

def create_cheat_sheet_docx(cheat_sheet_data, font="Calibri", font_size=11):
    doc = Document() #word doc

    # page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    try:
        #go through each section of the structured response
        for section_name, content in cheat_sheet_data.items():
            # create a heading for the section
            section = doc.add_heading(clean_text(section_name), level=1)
            section.style.font.name = font
            section.style.font.size = Pt(font_size + 2)
            section.paragraph_format.space_before = Pt(3)
            section.paragraph_format.space_after = Pt(3)

            #add content from dictionary
            if isinstance(content, dict):
                for subsection, details in content.items():
                    subsection_para = doc.add_paragraph()
                    run = subsection_para.add_run(clean_text(subsection))
                    run.bold = True
                    run.font.size = Pt(font_size)
                    
                    #check if the details are code or text
                    if isinstance(details, dict):
                        for key, value in details.items():
                            if key == "Code":
                                # Format the code section
                                code_para = doc.add_paragraph()
                                code_run = code_para.add_run(clean_text(value))
                                code_run.font.name = 'Courier New'
                                code_run.font.size = Pt(font_size + 1)
                                code_run.font.color.rgb = RGBColor(45, 34, 143)  # Gray color
                                code_para.paragraph_format.left_indent = Inches(0.5)
                            else:
                                detail_para = doc.add_paragraph(style='List Bullet')
                                for word in clean_text(key).split():
                                    word_run = detail_para.add_run(word + " ")
                                detail_para.add_run(": ")
                                for word in clean_text(value).split():
                                    word_run = detail_para.add_run(word + " ")
                                detail_para.paragraph_format.space_before = Pt(3)
                                detail_para.paragraph_format.space_after = Pt(3)
                    else:
                        detail_para = doc.add_paragraph(style='List Bullet')
                        for word in clean_text(details).split():
                            word_run = detail_para.add_run(word + " ")
                        detail_para.paragraph_format.space_before = Pt(3)
                        detail_para.paragraph_format.space_after = Pt(3)
            else:
                para = doc.add_paragraph(clean_text(content))
                para.style.font.size = Pt(font_size)
                para.paragraph_format.space_before = Pt(3)
                para.paragraph_format.space_after = Pt(3)

    except Exception as e:
        print(f"Document creation error: {str(e)}")
        raise

    # save doc
    doc_buffer = BytesIO()
    doc.save(doc_buffer)
    doc_buffer.seek(0)
    return doc_buffer



if "append_counter" not in st.session_state:
    st.session_state.append_counter = 0

def get_content_label(section):
    #add label to the appended sections (based on what section they were appended from)
    section = section.lower().strip()
    
    # check what type of content is being added (for the header)
    if section.startswith('# ') or section.startswith('## '):
        return "Definition"
    elif any(term in section for term in ['formula', 'equation', '\\frac', '\\sigma']):
        return "Formula"
    elif section.startswith('```'):
        return "Code Example"
    elif section.startswith(('1.', '2.', '3.')):
        return "Key Point"
    elif 'example' in section:
        return "Example"
    else:
        return "Additional Info"

def add_to_cheat_sheet(section, content_type):
    # add the selected content to the cheat sheet.
    try:
        st.session_state.append_counter += 1
        
        label = get_content_label(section)
        section = section.replace('**', '')
        
        
        new_content = {
            f"{label} {st.session_state.append_counter}": {
                label: section.strip() + "\n"
            }
        }
        
        update_cheat_sheet(
            cheat_sheet_format=st.session_state.cheat_sheet_format,
            question="",
            structured_response=new_content
        )
        
        return True
    except Exception as e:
        st.error(f"Error adding to cheat sheet: {str(e)}")
        return False

def update_cheat_sheet(cheat_sheet_format, question, structured_response):
    def deep_merge(target, source):
        for key, value in source.items():
            if key in target:
                if isinstance(value, dict) and isinstance(target[key], dict):
                    deep_merge(target[key], value)
                elif isinstance(value, str) and isinstance(target[key], dict):
                    continue
                elif isinstance(value, dict) and isinstance(target[key], str):
                    target[key] = {"General": target[key]}
                    deep_merge(target[key], value)
                else:
                    if value not in target[key]:
                        target[key] = f"{target[key]}\n\n{value}"
            else:
                target[key] = value

    try:
        if structured_response:
            deep_merge(cheat_sheet_format, structured_response)
    except Exception as e:
        st.error(f"Error updating cheat sheet: {str(e)}")
        print(f"Update error: {str(e)}")

def display_cheat_sheet(cheat_sheet_data):
    st.sidebar.title("Cheat Sheet")
    
    if not cheat_sheet_data:
        st.sidebar.write("No cheat sheet yet. Start asking questions!")
        return

    try:
        for section_name, topics in cheat_sheet_data.items():
            st.sidebar.markdown(f"### {clean_text(section_name)}")
            
            if isinstance(topics, dict):
                for topic_name, content in topics.items():
                    with st.sidebar.expander(clean_text(topic_name)):
                        if isinstance(content, dict):
                            for subtopic, text in content.items():
                                st.markdown(f"**{clean_text(subtopic)}:** {clean_text(text)}")
                        else:
                            st.markdown(clean_text(content))
            else:
                st.sidebar.markdown(clean_text(topics))

        # download the cheat sheet
        with st.spinner("Preparing download..."):
            try:
                cheat_sheet_file = create_cheat_sheet_docx(cheat_sheet_data)
                # Use current title directly for the filename
                file_name = f"{st.session_state.current_title}.docx"
                st.sidebar.download_button(
                    label="Download Cheat Sheet",
                    data=cheat_sheet_file,
                    file_name=file_name,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    on_click=st.cache_resource
                )
            except Exception as e:
                st.sidebar.error(f"Error generating cheat sheet: {e}")
                print(f"Document generation error: {str(e)}")


    except Exception as e:
        st.sidebar.error(f"Error displaying cheat sheet: {str(e)}")
        print(f"Display error: {str(e)}")