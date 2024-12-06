from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Twips

class DocumentFormatter:
    def __init__(self, font="Calibri", base_font_size=11):
        self.font = font
        self.base_font_size = base_font_size
        self.colors = {
            'heading': RGBColor(0, 51, 102),  # Dark blue
            'code': RGBColor(61, 61, 61),     # Dark gray
            'equation': RGBColor(0, 0, 0)     # Black
        }

    def format_equation(self, doc, equation_text):
        """Format mathematical equations with proper styling"""
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        # Replace common math symbols with proper Unicode characters
        math_replacements = {
            'inf': '∞',
            '<=': '≤',
            '>=': '≥',
            '!=': '≠',
            '==': '=',
            '->': '→',
            '<-': '←',
            'sqrt': '√',
            'alpha': 'α',
            'beta': 'β',
            'theta': 'θ',
            'pi': 'π',
            'sum': 'Σ',
            '...': '…'
        }
        
        for old, new in math_replacements.items():
            equation_text = equation_text.replace(old, new)
        
        run = para.add_run(equation_text)
        run.font.name = 'Cambria Math'
        run.font.size = Pt(self.base_font_size)
        return para

    def add_table(self, doc, headers, rows):
        """Add a formatted table"""
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        table.autofit = True
        
        # Add headers
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(self.base_font_size)
                    run.font.name = self.font
        
        # Add data rows
        for row_data in rows:
            cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                cells[i].text = str(cell_data)
                for paragraph in cells[i].paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(self.base_font_size)
                        run.font.name = self.font
        
        doc.add_paragraph()  # Add spacing after table
        return table

    def format_code(self, doc, code_text, language=""):
        """Format code blocks with proper styling"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.5)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(6)
        
        if language:
            lang_run = para.add_run(f"{language}:\n")
            lang_run.font.bold = True
            lang_run.font.size = Pt(self.base_font_size - 1)
        
        code_run = para.add_run(code_text)
        code_run.font.name = 'Consolas'
        code_run.font.size = Pt(self.base_font_size - 1)
        code_run.font.color.rgb = self.colors['code']
        return para

    def add_definition(self, doc, term, definition):
        """Add a term and its definition"""
        para = doc.add_paragraph()
        term_run = para.add_run(f"{term}: ")
        term_run.bold = True
        term_run.font.size = Pt(self.base_font_size)
        
        def_run = para.add_run(definition)
        def_run.font.size = Pt(self.base_font_size)
        return para

    def add_bullet_point(self, doc, text, level=0):
        """Add a bullet point with proper indentation"""
        para = doc.add_paragraph(style='List Bullet')
        para.paragraph_format.left_indent = Inches(0.25 * (level + 1))
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        run.font.size = Pt(self.base_font_size)
        return para

    def set_document_style(self, doc):
        """Set up the basic document styling"""
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def format_heading(self, doc, text, level):
        """Format headings with proper styling"""
        heading = doc.add_heading(text, level=level)
        
        # Adjust font size based on heading level
        sizes = {0: 16, 1: 14, 2: 12, 3: 11}
        font_size = sizes.get(level, self.base_font_size)
        
        for run in heading.runs:
            run.font.name = self.font
            run.font.size = Pt(font_size)
            if level <= 2:  # Only color main headings
                run.font.color.rgb = self.colors['heading']
        
        # Add appropriate spacing
        if level == 0:  # Title
            heading.paragraph_format.space_after = Pt(12)
        else:
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
        
        return heading

    def add_note(self, doc, text):
        """Add a note or important point"""
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Inches(0.25)
        note_run = para.add_run("Note: ")
        note_run.bold = True
        note_run.font.size = Pt(self.base_font_size - 1)
        
        text_run = para.add_run(text)
        text_run.font.italic = True
        text_run.font.size = Pt(self.base_font_size - 1)
        return para